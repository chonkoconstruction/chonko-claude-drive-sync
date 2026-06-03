#!/usr/bin/env python3
"""
update_drive.py — Chonko Drive Sync Worker

Converts markdown content to a properly formatted .docx (headings, tables,
bold, bullets) using python-docx, then uploads it to Google Drive via the
Drive API. Drive auto-converts it back to a Google Doc with full formatting.
"""
import os, sys, json, re
from io import BytesIO
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaInMemoryUpload
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def log(msg=""): print(msg, flush=True)

# ── Markdown parser ────────────────────────────────────────────────────────────

def is_table_line(line):
    return bool(line.strip()) and line.strip().startswith('|')

def is_separator_line(line):
    return bool(re.match(r'^\s*\|[-|\s:]+\|\s*$', line))

def parse_table(table_lines):
    """Parse markdown table lines into a list of rows (list of cell strings)."""
    rows = []
    for line in table_lines:
        if is_separator_line(line):
            continue
        cells = [c.strip() for c in re.split(r'(?<!\\)\|', line.strip('|'))]
        if cells:
            rows.append(cells)
    return rows

def add_formatted_run(paragraph, text):
    """Add text with inline **bold** and *italic* support."""
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
    for part in re.split(pattern, text):
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        elif part:
            paragraph.add_run(part)

def md_to_docx(text):
    doc = Document()

    # Standard US Letter portrait with 1-inch margins
    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)

    lines = text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── Table block ───────────────────────────────────────────────────────
        if is_table_line(line):
            table_lines = []
            while i < len(lines) and is_table_line(lines[i]):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table(table_lines)
            if rows:
                num_cols = max(len(r) for r in rows)
                try:
                    tbl = doc.add_table(rows=len(rows), cols=num_cols, style='Table Grid')
                except KeyError:
                    tbl = doc.add_table(rows=len(rows), cols=num_cols)
                for r_idx, row_data in enumerate(rows):
                    for c_idx in range(num_cols):
                        cell_text = row_data[c_idx] if c_idx < len(row_data) else ''
                        cell = tbl.cell(r_idx, c_idx)
                        cell.paragraphs[0].clear()
                        run = cell.paragraphs[0].add_run(cell_text)
                        if r_idx == 0:
                            run.bold = True
            doc.add_paragraph()
            continue

        # ── Headings ──────────────────────────────────────────────────────────
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)

        # ── Horizontal rule ───────────────────────────────────────────────────
        elif re.match(r'^[-_*]{3,}\s*$', line):
            doc.add_paragraph('─' * 60)

        # ── Bullet lists ──────────────────────────────────────────────────────
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_run(p, line[2:])

        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            add_formatted_run(p, re.sub(r'^\d+\.\s', '', line))

        # ── Empty line ────────────────────────────────────────────────────────
        elif line.strip() == '':
            doc.add_paragraph()

        # ── Normal paragraph ──────────────────────────────────────────────────
        else:
            p = doc.add_paragraph()
            add_formatted_run(p, line)

        i += 1

    return doc

# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    try:
        sa_json = os.environ["GOOGLE_SERVICE_ACCOUNT_JSON"]
        file_id = os.environ["FILE_ID"]
        pf      = os.environ["PENDING_FILE"]
    except KeyError as e:
        sys.exit(f"Missing env var: {e}")

    doc_name = os.environ.get("DOC_NAME", "document")
    log("=== Chonko Drive Sync ===")
    log(f"Target : {doc_name}")
    log(f"File ID: {file_id}")

    with open(pf, "r", encoding="utf-8") as f:
        content = f.read()
    log(f"Content: {len(content):,} chars — first line: {content.splitlines()[0]}")

    log("Converting markdown to .docx…")
    doc = md_to_docx(content)
    buf = BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()
    log(f"DOCX size: {len(docx_bytes):,} bytes")

    creds = service_account.Credentials.from_service_account_info(
        json.loads(sa_json),
        scopes=["https://www.googleapis.com/auth/drive"]
    )
    drive = build("drive", "v3", credentials=creds, cache_discovery=False)

    log("Uploading to Drive…")
    media = MediaInMemoryUpload(
        docx_bytes,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    result = drive.files().update(fileId=file_id, media_body=media).execute()
    log(f"✅ Updated: {result.get('name')} ({result.get('id')})")

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        sys.exit(f"ERROR: {e}")
